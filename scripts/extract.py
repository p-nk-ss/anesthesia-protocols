# -*- coding: utf-8 -*-
"""Extract text from doc/ sources into doc/_extracted/*.txt (UTF-8).
Avoids passing Cyrillic filenames on the command line by listing the folder."""
import os, sys, io

DOC = os.path.join(os.path.dirname(__file__), '..', 'doc')
OUT = os.path.join(DOC, '_extracted')
os.makedirs(OUT, exist_ok=True)


def extract_docx(path):
    import docx
    d = docx.Document(path)
    lines = []
    # Walk body in order: paragraphs and tables
    from docx.oxml.ns import qn
    body = d.element.body
    # Simple sequential: paragraphs then tables (python-docx limitation).
    for p in d.paragraphs:
        if p.text.strip():
            style = (p.style.name or '') if p.style else ''
            prefix = ''
            if 'Heading 1' in style: prefix = '# '
            elif 'Heading 2' in style: prefix = '## '
            elif 'Heading 3' in style: prefix = '### '
            elif 'Heading' in style: prefix = '#### '
            lines.append(prefix + p.text)
    for ti, t in enumerate(d.tables):
        lines.append(f'\n[TABLE {ti+1}]')
        for row in t.rows:
            cells = [c.text.replace('\n', ' ').strip() for c in row.cells]
            lines.append(' | '.join(cells))
    return '\n'.join(lines)


def extract_pdf(path):
    import fitz
    doc = fitz.open(path)
    out = []
    bad = 0
    total = 0
    for i, page in enumerate(doc):
        t = page.get_text()
        out.append(f'===== PAGE {i+1} =====\n{t}')
        for ch in t:
            total += 1
            if ch == '�':
                bad += 1
    ratio = (bad / total) if total else 0
    return '\n'.join(out), ratio


def main():
    files = sorted(os.listdir(DOC))
    report = []
    for f in files:
        path = os.path.join(DOC, f)
        if not os.path.isfile(path):
            continue
        ext = f.lower().rsplit('.', 1)[-1]
        base = f.rsplit('.', 1)[0]
        try:
            if ext == 'docx':
                text = extract_docx(path)
                ratio = 0.0
            elif ext == 'pdf':
                text, ratio = extract_pdf(path)
            else:
                continue
        except Exception as e:
            report.append(f'{f}: ERROR {e}')
            continue
        outp = os.path.join(OUT, base + '.txt')
        with io.open(outp, 'w', encoding='utf-8') as fh:
            fh.write(text)
        report.append(f'{f}: {len(text)} chars, replacement_ratio={ratio:.3f}')
    with io.open(os.path.join(OUT, '_report.txt'), 'w', encoding='utf-8') as fh:
        fh.write('\n'.join(report))
    # print ascii-safe report
    for line in report:
        sys.stdout.buffer.write((line + '\n').encode('ascii', 'replace'))


if __name__ == '__main__':
    main()
